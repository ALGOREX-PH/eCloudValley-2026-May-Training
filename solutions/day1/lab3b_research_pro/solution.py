"""
Lab 3B — Reference solution.

Research assistant with preferred-source steering + md/docx/pdf output.
"""

import argparse
import re
from pathlib import Path

from agno.agent import Agent
from agno.models.openai import OpenAIChat
from agno.tools import tool
from agno.tools.duckduckgo import DuckDuckGoTools
from dotenv import load_dotenv

load_dotenv()

NOTES_DIR = Path("notes")

DEFAULT_PREFERRED_SOURCES = [
    "arxiv.org",
    "anthropic.com",
    "openai.com",
    "docs.agno.com",
    "research.google",
    "deepmind.google",
    "huggingface.co",
]


def _slugify(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^\w\s-]", "", value)
    value = re.sub(r"[\s_]+", "-", value)
    return value[:60] or "untitled"


def _to_latin1(text: str) -> str:
    """fpdf2 default font is latin-1; replace what it can't render."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _save_md(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _save_docx(path: Path, content: str) -> None:
    from docx import Document

    doc = Document()
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(str(path))


def _save_pdf(path: Path, content: str) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    for raw in content.splitlines():
        line = _to_latin1(raw.rstrip())
        if not line:
            pdf.ln(4)
        elif line.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 9, line[2:])
            pdf.set_font("Helvetica", size=11)
        elif line.startswith("## "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, line[3:])
            pdf.set_font("Helvetica", size=11)
        elif line.startswith("- "):
            pdf.multi_cell(0, 6, f"  • {line[2:]}")
        else:
            pdf.multi_cell(0, 6, line)

    pdf.output(str(path))


SAVERS = {
    "md": (".md", _save_md),
    "docx": (".docx", _save_docx),
    "pdf": (".pdf", _save_pdf),
}


@tool
def save_note(filename: str, content: str, format: str = "md") -> str:
    """Save a research brief to disk in the chosen format.

    Use this tool ONCE at the end of your research, after you have synthesized
    findings into a final markdown brief.

    Args:
        filename: A short, descriptive slug. Lowercase, hyphenated, no
            extension. Example: "agentic-ai-papers-may-2026".
        content: The full markdown content of the brief. Must include
            `# Title`, `## Summary`, `## Key Findings`, `## Sources`.
        format: One of "md", "docx", "pdf". Default "md".

    Returns:
        The absolute path of the file that was written.
    """
    if format not in SAVERS:
        raise ValueError(f"Unknown format: {format!r}. Use one of {list(SAVERS)}.")

    NOTES_DIR.mkdir(exist_ok=True)
    safe = _slugify(filename)
    suffix, saver = SAVERS[format]
    path = NOTES_DIR / f"{safe}{suffix}"
    saver(path, content)
    return str(path.resolve())


def build_agent(format: str, preferred_sites: list[str]) -> Agent:
    sites_clause = " OR ".join(f"site:{d}" for d in preferred_sites)

    return Agent(
        model=OpenAIChat(id="gpt-4o-mini"),
        tools=[DuckDuckGoTools(), save_note],
        markdown=True,
        tool_call_limit=8,
        instructions=[
            "You are a thorough research assistant.",
            "When given a topic, search the web 2–3 times with focused queries.",
            "Append this site filter to EACH search query so results come from preferred domains:",
            f"  ({sites_clause})",
            "Spread your queries across at least 3 of those domains when possible.",
            "If a preferred-source query returns nothing useful, you MAY drop the filter "
            "for ONE follow-up search — but mention in the brief which claims came from outside the preferred list.",
            "After at most 3 searches, STOP searching and write the brief.",
            "The brief must follow this exact structure:",
            "  # <Title>",
            "  ## Summary  (2–3 sentences)",
            "  ## Key Findings  (bullet list, each with a URL)",
            "  ## Sources  (numbered list of URLs)",
            "Every factual claim must include a URL.",
            f"At the end, call save_note exactly once with format={format!r}.",
            "Then reply with a one-line confirmation including the file path.",
        ],
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Research assistant with preferred sources + multi-format output",
    )
    parser.add_argument("topic", nargs="+", help="Research topic")
    parser.add_argument(
        "--format",
        choices=["md", "docx", "pdf"],
        default="md",
        help="Output format (default: md)",
    )
    parser.add_argument(
        "--sources",
        nargs="+",
        default=DEFAULT_PREFERRED_SOURCES,
        help="Preferred source domains",
    )
    args = parser.parse_args()

    topic = " ".join(args.topic)
    print(f"📚 Researching: {topic}")
    print(f"📄 Format:      {args.format}")
    print(f"🌐 Preferring:  {', '.join(args.sources)}\n")

    agent = build_agent(args.format, args.sources)
    agent.print_response(
        f"Research the topic: '{topic}'. Synthesize a brief and save it.",
        stream=True,
    )


if __name__ == "__main__":
    main()
